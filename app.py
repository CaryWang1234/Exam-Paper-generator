#!/usr/bin/env python3
# -*- coding: utf-8 -*-
import os
import re
import sys
import glob
import time
import uuid
import locale
import logging
import tempfile
import sqlite3
import configparser
from io import BytesIO, StringIO
from datetime import datetime

from flask import Flask, render_template, request, redirect, url_for, send_file, flash, jsonify, abort
from openai import OpenAI
from docx import Document
from docx.shared import Pt
from dotenv import load_dotenv
import markdown
from flask_mail import Mail, Message
from apscheduler.schedulers.background import BackgroundScheduler

# Windows 控制台 UTF-8 输出
if sys.platform.startswith('win'):
    os.environ['PYTHONIOENCODING'] = 'utf-8'
    if hasattr(sys.stdout, 'reconfigure'):
        sys.stdout.reconfigure(encoding='utf-8')
    if hasattr(sys.stderr, 'reconfigure'):
        sys.stderr.reconfigure(encoding='utf-8')

try:
    locale.setlocale(locale.LC_ALL, 'zh_CN.UTF-8')
except locale.Error:
    try:
        locale.setlocale(locale.LC_ALL, 'Chinese_China.UTF-8')
    except locale.Error:
        pass

# 路径常量（支持环境变量覆盖，便于测试隔离）
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
CONFIG_FILE = os.environ.get('APP_CONFIG_FILE', os.path.join(BASE_DIR, 'config.ini'))
DB_PATH = os.environ.get('STUDENTS_DB', os.path.join(BASE_DIR, 'students.db'))
LOG_FILE = os.path.join(BASE_DIR, 'log.txt')

# 日志系统：控制台 + log.txt（UTF-8）
logger = logging.getLogger('exam_generator')
logger.setLevel(logging.INFO)
if not logger.handlers:
    _formatter = logging.Formatter('%(asctime)s [%(levelname)s] %(message)s')
    _console_handler = logging.StreamHandler()
    _console_handler.setFormatter(_formatter)
    logger.addHandler(_console_handler)
    try:
        _file_handler = logging.FileHandler(LOG_FILE, encoding='utf-8')
        _file_handler.setFormatter(_formatter)
        logger.addHandler(_file_handler)
    except OSError:
        logger.warning('无法创建日志文件 log.txt，仅输出到控制台')

# 加载环境变量
load_dotenv()

app = Flask(__name__)

# 配置文件管理
DEFAULT_CONFIG = {
    'mail_server': 'smtp.gmail.com',
    'mail_port': '587',
    'mail_use_tls': 'True',
    'mail_username': '',
    'mail_password': '',
    'mail_default_sender': '',
    'deepseek_api_key': '',
    'deepseek_base_url': 'https://api.deepseek.com',
    'deepseek_model': 'deepseek-chat',
    'daily_practice_time': '07:00',
}

def load_config():
    """加载配置文件，不存在时创建默认配置"""
    config = configparser.ConfigParser()

    if not os.path.exists(CONFIG_FILE):
        config['DEFAULT'] = dict(DEFAULT_CONFIG)
        with open(CONFIG_FILE, 'w', encoding='utf-8') as f:
            config.write(f)

    config.read(CONFIG_FILE, encoding='utf-8')
    return config

# 加载配置
config = load_config()

# secret_key 持久化：首次运行生成并写入配置文件，之后复用
_secret_key = config.get('DEFAULT', 'secret_key', fallback='')
if not _secret_key:
    _secret_key = os.urandom(24).hex()
    config['DEFAULT']['secret_key'] = _secret_key
    try:
        with open(CONFIG_FILE, 'w', encoding='utf-8') as f:
            config.write(f)
    except OSError as e:
        logger.warning(f'secret_key 写入配置文件失败: {e}')
app.secret_key = _secret_key

# 邮件配置
app.config['MAIL_SERVER'] = config.get('DEFAULT', 'mail_server', fallback='smtp.gmail.com')
app.config['MAIL_PORT'] = config.getint('DEFAULT', 'mail_port', fallback=587)
app.config['MAIL_USE_TLS'] = config.getboolean('DEFAULT', 'mail_use_tls', fallback=True)
app.config['MAIL_USERNAME'] = config.get('DEFAULT', 'mail_username', fallback='')
app.config['MAIL_PASSWORD'] = config.get('DEFAULT', 'mail_password', fallback='')
app.config['MAIL_DEFAULT_SENDER'] = config.get('DEFAULT', 'mail_default_sender', fallback='')

mail = Mail(app)

# DeepSeek API 配置（惰性初始化，密钥缺失时不阻止应用启动）
def get_ai_config():
    """读取最新的 AI 配置，环境变量 DEEPSEEK_API_KEY 优先于配置文件"""
    cfg = load_config()
    api_key = os.environ.get('DEEPSEEK_API_KEY') or cfg.get('DEFAULT', 'deepseek_api_key', fallback='')
    base_url = cfg.get('DEFAULT', 'deepseek_base_url', fallback='') or 'https://api.deepseek.com'
    model = cfg.get('DEFAULT', 'deepseek_model', fallback='deepseek-chat')
    return api_key, base_url, model

_ai_client = None
_ai_client_key = None

def get_ai_client():
    """惰性获取 OpenAI 客户端，未配置密钥时返回 None"""
    global _ai_client, _ai_client_key
    api_key, base_url, _ = get_ai_config()
    if not api_key:
        return None
    if _ai_client is None or _ai_client_key != api_key:
        _ai_client = OpenAI(api_key=api_key, base_url=base_url)
        _ai_client_key = api_key
    return _ai_client

if not get_ai_config()[0]:
    logger.warning('DeepSeek API 密钥未配置，生成功能暂不可用。请在系统配置页面或 .env 文件中设置。')

# 邮箱格式校验
EMAIL_RE = re.compile(r'^[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}$')

def is_valid_email(email):
    """校验邮箱格式"""
    return bool(email) and bool(EMAIL_RE.match(email.strip()))

# 数据库初始化（改进UTF-8支持）
def init_db():
    """初始化数据库，确保UTF-8编码支持"""
    conn = sqlite3.connect(DB_PATH,
                           detect_types=sqlite3.PARSE_DECLTYPES,
                           check_same_thread=False)
    conn.execute("PRAGMA encoding = 'UTF-8'")
    conn.execute("PRAGMA foreign_keys = ON")

    c = conn.cursor()

    # 学生订阅表（添加详细信息字段）
    c.execute('''CREATE TABLE IF NOT EXISTS student_subscriptions
                 (id INTEGER PRIMARY KEY AUTOINCREMENT,
                  email TEXT UNIQUE NOT NULL,
                  name TEXT,
                  grade INTEGER DEFAULT 7,
                  subject TEXT,
                  knowledge_scope TEXT,
                  learning_goals TEXT,
                  special_requirements TEXT,
                  difficulty TEXT DEFAULT '中等',
                  daily_questions INTEGER DEFAULT 5,
                  preferred_time TEXT DEFAULT '07:00',
                  subscription_date TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                  is_active BOOLEAN DEFAULT 1)''')

    # 练习记录表
    c.execute('''CREATE TABLE IF NOT EXISTS practice_records
                 (id INTEGER PRIMARY KEY AUTOINCREMENT,
                  email TEXT NOT NULL,
                  practice_date DATE NOT NULL,
                  questions TEXT,
                  score INTEGER,
                  completed BOOLEAN DEFAULT 0,
                  feedback TEXT,
                  FOREIGN KEY (email) REFERENCES student_subscriptions (email))''')

    # 升级旧版表结构，补齐缺失的列
    _migrate_db(conn)

    conn.commit()
    conn.close()

def _migrate_db(conn):
    """升级旧版数据库表结构，补齐缺失的列"""
    subscription_columns = {
        'grade': 'INTEGER DEFAULT 7',
        'knowledge_scope': 'TEXT',
        'learning_goals': 'TEXT',
        'special_requirements': 'TEXT',
        'preferred_time': "TEXT DEFAULT '07:00'",
    }
    existing = {row[1] for row in conn.execute('PRAGMA table_info(student_subscriptions)')}
    for column, definition in subscription_columns.items():
        if column not in existing:
            conn.execute(f'ALTER TABLE student_subscriptions ADD COLUMN {column} {definition}')
            logger.info(f'数据库迁移：student_subscriptions 新增列 {column}')

    record_columns = {'feedback': 'TEXT'}
    existing_records = {row[1] for row in conn.execute('PRAGMA table_info(practice_records)')}
    for column, definition in record_columns.items():
        if column not in existing_records:
            conn.execute(f'ALTER TABLE practice_records ADD COLUMN {column} {definition}')
            logger.info(f'数据库迁移：practice_records 新增列 {column}')
    conn.commit()

# 改进的数据库连接函数
def get_db_connection():
    """获取数据库连接，确保UTF-8支持"""
    conn = sqlite3.connect(DB_PATH,
                           detect_types=sqlite3.PARSE_DECLTYPES,
                           check_same_thread=False)
    conn.execute("PRAGMA encoding = 'UTF-8'")
    conn.row_factory = sqlite3.Row
    return conn

# 初始化数据库并升级旧版表结构
init_db()

def generate_daily_practice(subject, difficulty, question_count=5, grade=None, knowledge_scope=None, learning_goals=None):
    """生成每日练习题（支持详细信息定制）"""
    client = get_ai_client()
    if client is None:
        return "生成练习题失败: 尚未配置 DeepSeek API 密钥，请先在系统配置页面设置。"

    _, _, model = get_ai_config()
    grade_info = f"{grade}年级" if grade else ""
    knowledge_info = f"，知识范围：{knowledge_scope}" if knowledge_scope else ""
    goals_info = f"，学习目标：{learning_goals}" if learning_goals else ""

    system_prompt = f"""你是一位专业的{subject}教师，负责为{grade_info}学生生成每日练习题{knowledge_info}{goals_info}。

    生成要求：
    1. 难度级别：{difficulty}
    2. 题目数量：{question_count}道
    3. 题目类型：选择题、填空题、简答题混合
    4. 题目内容要符合学生的认知水平和学习进度
    5. 包含详细的答案解析和学习建议
    6. 使用Markdown格式输出

    输出格式：
    # {subject}每日练习题 - {datetime.now().strftime('%Y年%m月%d日')}

    ## 今日学习目标
    [根据学习目标编写]

    ## 练习题

    ### 选择题
    1. 题目...
       A. 选项A
       B. 选项B
       C. 选项C
       D. 选项D
       **答案：** A
       **解析：** 详细解析...
       **学习建议：** 相关学习建议...

    ### 填空题
    2. 题目...
       **答案：** 正确答案
       **解析：** 详细解析...
       **学习建议：** 相关学习建议...

    ### 简答题
    3. 题目...
       **参考答案：** 详细解答
       **学习建议：** 相关学习建议...

    ## 今日总结
    [根据练习内容编写总结]
    """

    user_prompt = f"""
    请为{grade_info if grade_info else ''}{subject}科目生成{difficulty}难度的每日练习题，共{question_count}道题目。
    {f'学生年级：{grade}年级，请根据该年级的认知水平出题。' if grade else ''}
    {f'重点知识范围：{knowledge_scope}' if knowledge_scope else ''}
    {f'学习目标：{learning_goals}' if learning_goals else ''}
题目应该：
    1. 覆盖该科目的重点知识点
    2. 难度适中，适合学生日常练习
    3. 包含详细的解析和学习建议
    4. 帮助学生巩固知识和提高能力
    """

    try:
        response = client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ],
            stream=False
        )
        return response.choices[0].message.content
    except Exception as e:
        logger.error(f"生成练习题失败: {e}")
        return f"生成练习题失败: {str(e)}"

def send_practice_for_time_slot(time_slot):
    """向偏好时间匹配的所有活跃学生发送每日练习题"""
    default_time = _normalize_time(config.get('DEFAULT', 'daily_practice_time', fallback='07:00'))
    conn = get_db_connection()

    students = conn.execute('''SELECT email, name, grade, subject, knowledge_scope, learning_goals,
                                      difficulty, daily_questions, preferred_time
                               FROM student_subscriptions
                               WHERE is_active = 1''').fetchall()

    for student in students:
        # 仅向偏好时间与当前任务时段匹配的学生发送
        if _normalize_time(student['preferred_time'], default_time) != time_slot:
            continue

        email = student['email']
        name = student['name']
        grade = student['grade']
        subject = student['subject']
        knowledge_scope = student['knowledge_scope']
        learning_goals = student['learning_goals']

        # 生成每日练习题（使用详细信息）
        practice_content = generate_daily_practice(
            subject, student['difficulty'], student['daily_questions'],
            grade, knowledge_scope, learning_goals
        )
        if practice_content.startswith('生成练习题失败'):
            logger.error(f"为 {email} 生成练习题失败: {practice_content}")
            continue
        practice_html = markdown.markdown(practice_content)

        # 发送邮件
        try:
            msg = Message(
                subject=f"{subject}每日练习题 - {datetime.now().strftime('%Y年%m月%d日')}",
                recipients=[email],
                html=f"""
                <html>
                <body>
                    <h2>亲爱的{name}同学：</h2>
                    <p>这是为您定制的{subject}每日练习题，请认真完成！</p>
                    {f'<p><strong>年级：</strong>{grade}年级</p>' if grade else ''}
                    {f'<p><strong>知识范围：</strong>{knowledge_scope}</p>' if knowledge_scope else ''}
                    {f'<p><strong>学习目标：</strong>{learning_goals}</p>' if learning_goals else ''}
                    <div style="background-color: #f8f9fa; padding: 20px; border-radius: 10px;">
                        {practice_html}
                    </div>
                    <p style="color: #666; font-size: 12px; margin-top: 20px;">
                        此邮件由智能试卷生成系统自动发送，如需取消订阅请访问系统设置。
                    </p>
                </body>
                </html>
                """
            )
            mail.send(msg)

            # 记录发送历史
            conn.execute('''INSERT INTO practice_records (email, practice_date, questions)
                            VALUES (?, ?, ?)''',
                         (email, datetime.now().date(), practice_content))
            conn.commit()

            logger.info(f"已发送个性化练习题给 {email}")
        except Exception as e:
            logger.error(f"发送邮件失败 {email}: {e}")

    conn.close()

# 定时任务调度器
scheduler = BackgroundScheduler()

def _normalize_time(value, fallback='07:00'):
    """规范化 HH:MM 时间格式，非法值回退到默认时间"""
    if value and re.match(r'^\d{1,2}:\d{2}$', value.strip()):
        hour, minute = value.strip().split(':')
        if 0 <= int(hour) <= 23 and 0 <= int(minute) <= 59:
            return f"{int(hour):02d}:{int(minute):02d}"
    return fallback

def ensure_time_slot(time_slot):
    """确保指定时段存在一个每日练习发送任务（幂等）"""
    job_id = f'daily_practice_{time_slot.replace(":", "")}'
    if scheduler.get_job(job_id) is None:
        hour, minute = map(int, time_slot.split(':'))
        scheduler.add_job(
            func=send_practice_for_time_slot,
            trigger='cron',
            hour=hour,
            minute=minute,
            args=[time_slot],
            id=job_id
        )
        logger.info(f"已注册每日练习发送任务：{time_slot}")

def schedule_daily_practice():
    """按全局默认时间与各活跃学生的偏好时间分别注册发送任务"""
    default_time = _normalize_time(config.get('DEFAULT', 'daily_practice_time', fallback='07:00'))
    ensure_time_slot(default_time)

    # 为每个活跃学生的偏好时间段注册任务
    conn = get_db_connection()
    rows = conn.execute(
        "SELECT DISTINCT preferred_time FROM student_subscriptions WHERE is_active = 1"
    ).fetchall()
    conn.close()
    for row in rows:
        ensure_time_slot(_normalize_time(row['preferred_time'], default_time))

    # 定期清理临时试卷文件
    scheduler.add_job(func=cleanup_temp_files, trigger='interval', hours=6, id='temp_cleanup')

    scheduler.start()

def cleanup_temp_files(max_age_hours=24):
    """清理临时目录中超过指定时长的试卷 .md 文件"""
    temp_dir = tempfile.gettempdir()
    now = time.time()
    removed = 0
    for path in glob.glob(os.path.join(temp_dir, '*.md')):
        try:
            if now - os.path.getmtime(path) > max_age_hours * 3600:
                os.remove(path)
                removed += 1
        except OSError:
            continue
    if removed:
        logger.info(f"已清理 {removed} 个过期临时试卷文件")

# 启动定时任务
try:
    schedule_daily_practice()
    logger.info(f"每日练习定时任务已启动，默认发送时间：{config.get('DEFAULT', 'daily_practice_time', fallback='07:00')}")
    cleanup_temp_files()
except Exception as e:
    logger.error(f"定时任务启动失败: {e}")

def generate_exam(material, exam_type, difficulty, total_score, time_limit, question_types, instructions):
    """调用DeepSeek API生成试卷"""
    client = get_ai_client()
    if client is None:
        return "生成试卷失败: 尚未配置 DeepSeek API 密钥，请先在系统配置页面设置。"

    _, _, model = get_ai_config()
    system_prompt = """你是一位专业的试卷生成专家。请根据提供的资料和配置参数生成一份高质量试卷。

    试卷生成要求：
    1. 根据试卷类型、难度级别、总分和考试时间合理分配题目
    2. 题目类型多样，包括选择题、填空题、简答题、应用题等
    3. 题目难度分布合理，符合指定的难度级别
    4. 包含清晰的题目描述、分值和参考答案
    5. 使用规范的Markdown格式输出

    输出格式要求：
    # 试卷名称
    ## 考试信息（类型、难度、总分、时间）
    ### 一、选择题（每题分值）
    1. 题目内容...
       A. 选项A
       B. 选项B
       C. 选项C
       D. 选项D
       **答案：** A

    ### 二、填空题
    1. 题目内容...
       **答案：** 正确答案

    ### 三、简答题
    1. 题目内容...
       **参考答案：** 详细解答

    ## 参考答案汇总
    """

    # 构建详细的用户提示
    user_prompt = f"""
    试卷生成配置：
    - 试卷类型：{exam_type}
    - 难度级别：{difficulty}
    - 总分：{total_score}分
    - 考试时间：{time_limit}分钟
    - 题目类型要求：{question_types if question_types else '自动分配'}
    - 自定义指令：{instructions if instructions else '无'}

    资料内容：
    {material}

    请根据以上配置和资料生成一份完整的试卷，严格按照Markdown格式输出。
    """

    try:
        response = client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ],
            stream=False
        )
        return response.choices[0].message.content
    except Exception as e:
        logger.error(f"API调用错误: {e}")
        return f"生成试卷失败: {str(e)}"

def extract_docx_text(file_stream):
    """从 DOCX 文件流中提取纯文本内容"""
    doc = Document(file_stream)
    return '\n'.join(p.text for p in doc.paragraphs if p.text.strip())

def _strip_bold_markers(text):
    """去除 Markdown 加粗标记"""
    return re.sub(r'\*\*(.+?)\*\*', r'\1', text)

def convert_markdown_to_docx(markdown_content):
    """将Markdown内容转换为DOCX文档"""
    doc = Document()

    # 设置文档样式
    styles = doc.styles
    title_style = styles['Heading 1']
    title_style.font.size = Pt(18)
    title_style.font.bold = True

    heading_style = styles['Heading 2']
    heading_style.font.size = Pt(14)
    heading_style.font.bold = True

    # 逐行处理Markdown内容，每行独立成段，避免内容串段
    for line in markdown_content.split('\n'):
        line = line.strip()
        if not line:
            continue

        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(_strip_bold_markers(line[2:]), style='List Bullet')
        elif re.match(r'^\d+[.、]\s*', line):
            doc.add_paragraph(_strip_bold_markers(re.sub(r'^\d+[.、]\s*', '', line)),
                              style='List Number')
        elif line.startswith('**答案') or line.startswith('**参考答案'):
            run = doc.add_paragraph().add_run(_strip_bold_markers(line))
            run.bold = True
        else:
            doc.add_paragraph(_strip_bold_markers(line))

    # 保存到内存
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def _register_chinese_font():
    """注册中文字体用于PDF输出，返回字体名；未找到可用字体时返回 None"""
    if sys.platform.startswith('win'):
        font_dir = os.path.join(os.environ.get('WINDIR', r'C:\Windows'), 'Fonts')
        candidates = [os.path.join(font_dir, name)
                      for name in ('simhei.ttf', 'msyh.ttf', 'simkai.ttf', 'simsun.ttc')]
    else:
        candidates = [
            '/usr/share/fonts/truetype/wqy/wqy-microhei.ttc',
            '/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc',
            '/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc',
            '/usr/share/fonts/truetype/droid/DroidSansFallbackFull.ttf',
        ]

    try:
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
    except ImportError:
        return None

    for path in candidates:
        if os.path.exists(path):
            try:
                pdfmetrics.registerFont(TTFont('ChineseFont', path))
                return 'ChineseFont'
            except Exception as e:
                logger.warning(f"中文字体注册失败 {path}: {e}")
    return None

def convert_markdown_to_pdf(markdown_content):
    """将Markdown内容转换为PDF（基于 xhtml2pdf，注册中文字体后支持中文）"""
    try:
        from xhtml2pdf import pisa
    except ImportError:
        logger.error("PDF 导出需要安装 xhtml2pdf: pip install xhtml2pdf")
        return None

    html_body = markdown.markdown(markdown_content)
    font_family = _register_chinese_font() or 'Helvetica'

    html_content = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <style>
        @page {{ size: A4; margin: 2cm; }}
        body {{ font-family: {font_family}; font-size: 12px; line-height: 1.6; }}
        h1 {{ font-size: 18px; color: #2c3e50; }}
        h2 {{ font-size: 15px; color: #34495e; }}
        h3 {{ font-size: 13px; color: #34495e; }}
    </style>
</head>
<body>
    {html_body}
</body>
</html>"""

    buffer = BytesIO()
    try:
        status = pisa.CreatePDF(StringIO(html_content), dest=buffer, encoding='utf-8')
        if status.err:
            logger.error('xhtml2pdf 转换过程中出现错误')
            return None
    except Exception as e:
        logger.error(f"PDF转换失败: {e}")
        return None
    buffer.seek(0)
    return buffer

# 仅允许本机访问的地址
LOCAL_ADDRESSES = {'127.0.0.1', '::1'}

def require_local_access():
    """限制管理接口仅本机可访问"""
    if request.remote_addr not in LOCAL_ADDRESSES:
        logger.warning(f"拒绝非本机访问管理接口: {request.remote_addr} -> {request.path}")
        abort(403)

@app.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        # 获取表单数据
        material_file = request.files.get('material')
        material_text = request.form.get('material_text', '')
        exam_type = request.form.get('exam_type', '考试')
        difficulty = request.form.get('difficulty', '中等')
        total_score = request.form.get('total_score', '100')
        time_limit = request.form.get('time_limit', '90')
        question_types = request.form.get('question_types', '')
        instructions = request.form.get('instructions', '')

        # 处理资料内容（改进UTF-8编码处理）
        material = ""
        if material_file and material_file.filename != '':
            filename = material_file.filename.lower()
            try:
                if filename.endswith('.txt') or filename.endswith('.md'):
                    # 尝试多种编码格式读取文件
                    file_content = material_file.read()
                    encodings = ['utf-8', 'gbk', 'gb2312', 'latin1']

                    for encoding in encodings:
                        try:
                            material = file_content.decode(encoding)
                            break
                        except UnicodeDecodeError:
                            continue
                    else:
                        # 如果所有编码都失败，使用utf-8并忽略错误
                        material = file_content.decode('utf-8', errors='ignore')
                elif filename.endswith('.docx'):
                    material = extract_docx_text(material_file)
                else:
                    flash("目前仅支持TXT、MD、DOCX格式文件")
                    return redirect(url_for('index'))
            except Exception as e:
                flash(f"文件读取失败: {str(e)}")
                return redirect(url_for('index'))

        # 如果没有上传文件，使用文本框内容
        if not material and material_text:
            material = material_text

        if not material:
            flash("请提供资料内容（上传文件或输入文本）")
            return redirect(url_for('index'))

        # 验证数值参数
        try:
            total_score = int(total_score)
            time_limit = int(time_limit)
            if total_score < 10 or total_score > 200:
                flash("总分应在10-200分之间")
                return redirect(url_for('index'))
            if time_limit < 10 or time_limit > 300:
                flash("考试时间应在10-300分钟之间")
                return redirect(url_for('index'))
        except ValueError:
            flash("请输入有效的数值参数")
            return redirect(url_for('index'))

        # 生成试卷
        exam_content = generate_exam(material, exam_type, difficulty, total_score, time_limit, question_types, instructions)
        if exam_content.startswith('生成试卷失败'):
            flash(exam_content)
            return redirect(url_for('index'))

        # 转换为HTML
        exam_html = markdown.markdown(exam_content)

        # 保存内容到临时文件（确保UTF-8编码）
        temp_dir = tempfile.gettempdir()
        file_id = str(uuid.uuid4())
        md_path = os.path.join(temp_dir, f"{file_id}.md")

        # 使用UTF-8编码保存文件，并添加BOM以支持Windows系统
        with open(md_path, 'w', encoding='utf-8-sig') as f:
            f.write(exam_content)

        return render_template('index.html',
                             exam_content=exam_html,
                             file_id=file_id)

    return render_template('index.html', exam_content=None, file_id=None)

@app.route('/student', methods=['GET', 'POST'])
def student_portal():
    """学生门户页面（支持详细信息输入）"""
    if request.method == 'POST':
        email = (request.form.get('email') or '').strip()
        name = request.form.get('name')
        grade = request.form.get('grade')
        subject = request.form.get('subject')
        knowledge_scope = request.form.get('knowledge_scope')
        learning_goals = request.form.get('learning_goals')
        special_requirements = request.form.get('special_requirements')
        difficulty = request.form.get('difficulty', '中等')
        daily_questions = request.form.get('daily_questions', '5')
        preferred_time = request.form.get('preferred_time', '07:00')
        action = request.form.get('action')

        # 邮箱格式校验
        if not is_valid_email(email):
            flash("邮箱格式不正确，请输入有效的邮箱地址")
            return redirect(url_for('student_portal'))

        # 数值参数校验
        try:
            daily_questions = int(daily_questions)
            if not 1 <= daily_questions <= 50:
                flash("每日题量应在1-50题之间")
                return redirect(url_for('student_portal'))
        except ValueError:
            flash("每日题量必须为数字")
            return redirect(url_for('student_portal'))

        # 年级转 int，空值存 NULL
        grade_value = int(grade) if grade and grade.strip().isdigit() else None

        conn = get_db_connection()
        c = conn.cursor()

        if action == 'subscribe':
            # 订阅每日练习（包含详细信息）
            try:
                c.execute('''INSERT OR REPLACE INTO student_subscriptions
                             (email, name, grade, subject, knowledge_scope, learning_goals,
                              special_requirements, difficulty, daily_questions, preferred_time, is_active)
                             VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, 1)''',
                         (email, name, grade_value, subject, knowledge_scope, learning_goals,
                          special_requirements, difficulty, daily_questions, preferred_time))
                conn.commit()
                # 为新订阅的偏好时间段动态注册发送任务
                default_time = _normalize_time(config.get('DEFAULT', 'daily_practice_time', fallback='07:00'))
                ensure_time_slot(_normalize_time(preferred_time, default_time))
                flash("订阅成功！您将每天收到个性化练习题。")
            except Exception as e:
                logger.error(f"订阅失败 {email}: {e}")
                flash(f"订阅失败: {str(e)}")

        elif action == 'unsubscribe':
            # 取消订阅
            c.execute("UPDATE student_subscriptions SET is_active = 0 WHERE email = ?", (email,))
            conn.commit()
            flash("已取消订阅每日练习服务")

        elif action == 'send_test':
            # 发送测试练习题（使用详细信息）
            practice_content = generate_daily_practice(
                subject, difficulty, daily_questions, grade_value, knowledge_scope, learning_goals
            )
            if practice_content.startswith('生成练习题失败'):
                flash(practice_content)
            else:
                practice_html = markdown.markdown(practice_content)
                try:
                    msg = Message(
                        subject=f"{subject}测试练习题",
                        recipients=[email],
                        html=f"""
                        <html>
                        <head>
                            <meta charset="UTF-8">
                        </head>
                        <body>
                            <h2>测试练习题</h2>
                            {f'<p><strong>年级：</strong>{grade_value}年级</p>' if grade_value else ''}
                            {f'<p><strong>知识范围：</strong>{knowledge_scope}</p>' if knowledge_scope else ''}
                            <div style="background-color: #f8f9fa; padding: 20px; border-radius: 10px;">
                                {practice_html}
                            </div>
                        </body>
                        </html>
                        """
                    )
                    mail.send(msg)
                    flash("测试练习题已发送到您的邮箱")
                except Exception as e:
                    logger.error(f"发送测试邮件失败 {email}: {e}")
                    flash(f"发送测试邮件失败: {str(e)}")

        conn.close()
        return redirect(url_for('student_portal'))

    return render_template('student.html')

@app.route('/admin/config', methods=['GET', 'POST'])
def admin_config():
    """系统配置管理页面（仅限本机访问）"""
    require_local_access()

    if request.method == 'POST':
        # 更新配置文件
        config['DEFAULT']['mail_server'] = request.form.get('mail_server', '')
        config['DEFAULT']['mail_port'] = request.form.get('mail_port', '587')
        config['DEFAULT']['mail_username'] = request.form.get('mail_username', '')
        config['DEFAULT']['mail_password'] = request.form.get('mail_password', '')
        config['DEFAULT']['mail_default_sender'] = request.form.get('mail_default_sender', '')
        config['DEFAULT']['deepseek_api_key'] = request.form.get('deepseek_api_key', '')
        config['DEFAULT']['daily_practice_time'] = request.form.get('daily_practice_time', '07:00')

        with open(CONFIG_FILE, 'w', encoding='utf-8') as f:
            config.write(f)

        logger.info("系统配置已更新")
        flash("配置已保存！邮件配置需要重启应用使配置生效，API密钥保存后立即生效。")
        return redirect(url_for('admin_config'))

    return render_template('admin_config.html', config=config['DEFAULT'])

@app.route('/api/send_practice', methods=['POST'])
def api_send_practice():
    """API接口：立即发送练习题（仅限本机访问）"""
    require_local_access()

    data = request.get_json(silent=True) or {}
    email = (data.get('email') or '').strip()
    subject = data.get('subject', '数学')
    difficulty = data.get('difficulty', '中等')
    question_count = data.get('question_count', 5)

    if not is_valid_email(email):
        return jsonify({'success': False, 'message': '邮箱格式不正确'}), 400

    practice_content = generate_daily_practice(subject, difficulty, question_count)
    if practice_content.startswith('生成练习题失败'):
        return jsonify({'success': False, 'message': practice_content})

    try:
        msg = Message(
            subject=f"{subject}即时练习题",
            recipients=[email],
            html=markdown.markdown(practice_content)
        )
        mail.send(msg)
        return jsonify({'success': True, 'message': '练习题已发送'})
    except Exception as e:
        logger.error(f"即时练习题发送失败 {email}: {e}")
        return jsonify({'success': False, 'message': str(e)})

@app.route('/download/<file_type>/<file_id>')
def download(file_type, file_id):
    temp_dir = tempfile.gettempdir()
    md_path = os.path.join(temp_dir, f"{file_id}.md")

    if not os.path.exists(md_path):
        flash("文件不存在或已过期")
        return redirect(url_for('index'))

    # 使用UTF-8编码读取文件
    with open(md_path, 'r', encoding='utf-8-sig') as f:
        content = f.read()

    if file_type == 'docx':
        buffer = convert_markdown_to_docx(content)
        return send_file(
            buffer,
            as_attachment=True,
            download_name=f'exam_paper_{datetime.now().strftime("%Y年%m月%d日")}.docx',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    elif file_type == 'html':
        html_content = markdown.markdown(content)
        html_doc = f'''<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta http-equiv="Content-Type" content="text/html; charset=UTF-8">
    <title>试卷 - {datetime.now().strftime("%Y年%m月%d日")}</title>
    <style>
        body {{ font-family: "Microsoft YaHei", "SimHei", sans-serif; line-height: 1.6; }}
        h1, h2, h3 {{ color: #2c3e50; }}
    </style>
</head>
<body>
    {html_content}
</body>
</html>'''
        buffer = BytesIO(html_doc.encode('utf-8'))
        return send_file(
            buffer,
            as_attachment=True,
            download_name=f'exam_paper_{datetime.now().strftime("%Y%m%d_%H%M")}.html',
            mimetype='text/html; charset=utf-8'
        )
    elif file_type == 'pdf':
        buffer = convert_markdown_to_pdf(content)
        if buffer is None:
            flash("PDF生成失败，请检查 xhtml2pdf 依赖是否安装，或改用HTML导出后打印为PDF")
            return redirect(url_for('index'))
        return send_file(
            buffer,
            as_attachment=True,
            download_name=f'exam_paper_{datetime.now().strftime("%Y%m%d_%H%M")}.pdf',
            mimetype='application/pdf'
        )
    else:
        flash("不支持的文件类型")
        return redirect(url_for('index'))

@app.route('/api/generate', methods=['POST'])
def api_generate():
    """API接口，用于前端异步调用"""
    data = request.get_json(silent=True) or {}

    material = data.get('material', '')
    exam_type = data.get('exam_type', '考试')
    difficulty = data.get('difficulty', '中等')
    total_score = data.get('total_score', 100)
    time_limit = data.get('time_limit', 90)
    question_types = data.get('question_types', '')
    instructions = data.get('instructions', '')

    if not material:
        return jsonify({'error': '请提供资料内容'}), 400

    exam_content = generate_exam(material, exam_type, difficulty, total_score, time_limit, question_types, instructions)
    if exam_content.startswith('生成试卷失败'):
        return jsonify({'error': exam_content}), 500

    return jsonify({
        'exam_content': exam_content,
        'exam_html': markdown.markdown(exam_content)
    })

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=False)
